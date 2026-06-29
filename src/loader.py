import tempfile
from langchain_community.document_loaders import PyPDFLoader
from langchain_core import document_loaders
from docx import Document as DocxDocument


def load_documents(uploaded_files):

    documents = []

    for file in uploaded_files:

        filename = file.name.lower()

        if filename.endswith(".pdf"):

            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
                tmp.write(file.read())

                loader = PyPDFLoader(tmp.name)

                pdf_docs = loader.load()

                for doc in pdf_docs:
                    doc.metadata["source"] = file.name

                documents.extend(pdf_docs)

        elif filename.endswith(".docx"):

            docx = DocxDocument(file)

            text = "\n".join(
                para.text for para in docx.paragraphs
            )

            documents.append(
                document_loaders.Document(
                    page_content=text,
                    metadata={
                        "source": file.name,
                        "page": 1
                    }
                )
            )

        elif filename.endswith(".txt"):

            text = file.read().decode("utf-8")

            documents.append(
                document_loaders.Document(
                    page_content=text,
                    metadata={
                        "source": file.name,
                        "page": 1
                    }
                )
            )

    return documents